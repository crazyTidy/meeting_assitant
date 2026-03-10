#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown转公文格式Word文档工具
将markdown文件转换为符合公文格式要求的Word文档
"""

import re
import sys
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.parser import OxmlElement
from docx.enum.style import WD_STYLE_TYPE


class MarkdownToDocxConverter:
    """Markdown转公文格式Word转换器"""

    # 字号映射（号数到磅值）
    FONT_SIZES = {
        '2号': 22,
        '3号': 16,
    }

    # 字体名称
    FONTS = {
        '方正小标宋简体': 'FZXiaoBiaoSong-B05S',
        '黑体': '黑体',
        '楷体GB2312': '楷体_GB2312',
        '仿宋GB2312': '仿宋_GB2312',
        'Times New Roman': 'Times New Roman',
    }

    def __init__(self, input_file=None, input_content=None, title="", output_folder=None):
        """初始化转换器

        Args:
            input_file: 输入的markdown文件路径
            input_content: 输入的markdown内容字符串
            title: 文档标题（当使用input_content时必须提供）
            output_folder: 输出文件夹路径（可选，默认为当前目录或文件所在目录）
        """
        if not input_file and not input_content:
            raise ValueError("必须提供 input_file 或 input_content 之一")

        self.input_file = None
        self.input_content = input_content

        if input_file:
            self.input_file = Path(input_file)
            if not self.input_file.exists():
                raise FileNotFoundError(f"文件不存在: {input_file}")

        # 获取当前日期（格式：YYYYMMDD）
        current_date = datetime.now().strftime('%Y%m%d')

        # 获取文件名（不含扩展名）
        file_stem = self.input_file.stem if self.input_file else title

        # 创建文件夹名称：YYYYMMDD_文件名
        folder_name = f"{current_date}_{file_stem}"

        # 创建文件夹路径
        if output_folder:
            self.output_folder = Path(output_folder) / folder_name
        elif self.input_file:
            self.output_folder = self.input_file.parent / folder_name
        else:
            self.output_folder = Path.cwd() / folder_name

        # 设置输出文件路径
        self.output_md = self.output_folder / f"{file_stem}.md"
        self.output_docx = self.output_folder / f"{file_stem}.docx"
        self.output_pdf = self.output_folder / f"{file_stem}.pdf"

        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]

        # 设置页面边距
        section.top_margin = Cm(3.7)      # 上边距37mm
        section.bottom_margin = Cm(3.5)   # 下边距35mm
        section.left_margin = Cm(2.8)     # 左边距28mm
        section.right_margin = Cm(2.6)    # 右边距26mm

        # 设置页面大小为A4
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    def _replace_quotes(self, text):
        """替换直引号为中文弯引号

        Args:
            text: 输入文本

        Returns:
            替换后的文本
        """
        # 使用正则表达式匹配成对的引号，类似Word的通配符方法
        # 第一步：将成对的引号替换为带临时标记的中文引号
        # 使用Unicode转义序列确保引号正确：\u201C是左双引号"，\u201D是右双引号"
        text = re.sub(r'"([^"]*)"', '■\u201C\\1\u201D■', text)

        # 第二步：删除临时标记
        text = text.replace('■', '')

        return text

    def _count_indent(self, line):
        """计算行首的缩进级别（每2个空格或1个tab为一级）

        Args:
            line: 文本行

        Returns:
            缩进级别（0-based）
        """
        indent = 0
        for char in line:
            if char == ' ':
                indent += 1
            elif char == '\t':
                indent += 2
            else:
                break
        return indent // 2

    def _is_table_separator(self, line):
        """判断是否为表格分隔行

        Args:
            line: 文本行

        Returns:
            是否为分隔行
        """
        stripped = line.strip()
        if not (stripped.startswith('|') and stripped.endswith('|')):
            return False
        cells = [c.strip() for c in stripped[1:-1].split('|')]
        # 所有单元格都是 '-' 或由 '-' 组成
        return all(re.match(r'^-+$', cell) for cell in cells if cell)

    def _parse_table(self, lines, start_index):
        """解析表格

        Args:
            lines: 所有行的列表
            start_index: 表格开始的索引

        Returns:
            (table_data, end_index) - 表格数据和结束索引
        """
        headers = None
        rows = []
        i = start_index

        # 解析表头
        if i < len(lines) and lines[i].strip().startswith('|'):
            header_line = lines[i].strip()
            headers = [c.strip() for c in header_line[1:-1].split('|')]
            i += 1

        # 跳过分隔行
        if i < len(lines) and self._is_table_separator(lines[i]):
            i += 1

        # 解析数据行
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('|'):
                break
            cells = [c.strip() for c in line[1:-1].split('|')]
            rows.append(cells)
            i += 1

        return {
            'type': 'table',
            'headers': headers,
            'rows': rows
        }, i - 1

    def _parse_markdown(self, content):
        """解析markdown内容

        Args:
            content: markdown文本内容

        Returns:
            解析后的结构化数据列表
        """
        lines = content.split('\n')
        parsed = []

        # 查找第一个标题的位置，跳过大模型的过渡性话语
        start_index = 0
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            # 检测是否为任何级别的标题
            if line_stripped.startswith('#') and ' ' in line_stripped:
                start_index = i
                break

        # 从第一个标题开始解析（如果没有标题，start_index为0，正常解析全部内容）
        i = start_index
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()

            # 跳过空行和分隔线
            if not line_stripped or line_stripped == '---':
                if parsed and parsed[-1]['type'] != 'empty':
                    parsed.append({'type': 'empty'})
                i += 1
                continue

            # 一级标题（文档标题）：# 标题
            if line.startswith('# '):
                parsed.append({
                    'type': 'title',
                    'text': line[2:].strip()
                })
            # 二级标题（一级标题）：## 一、...
            elif line.startswith('## '):
                parsed.append({
                    'type': 'heading1',
                    'text': line[3:].strip()
                })
            # 三级标题（二级标题）：### （一）...
            elif line.startswith('### '):
                parsed.append({
                    'type': 'heading2',
                    'text': line[4:].strip()
                })
            # 四级标题（三级标题）：#### 1. ...
            elif line.startswith('#### '):
                parsed.append({
                    'type': 'heading3',
                    'text': line[5:].strip()
                })
            # 列表项：- item 或 * item
            elif re.match(r'^(\s*)[-*]\s+', line):
                indent_level = self._count_indent(line)
                list_text = re.sub(r'^\s*[-*]\s+', '', line)
                parsed.append({
                    'type': 'list_item',
                    'text': list_text,
                    'level': indent_level
                })
            # 表格
            elif line_stripped.startswith('|') and line_stripped.endswith('|'):
                table_data, new_index = self._parse_table(lines, i)
                parsed.append(table_data)
                i = new_index
            # 普通段落
            else:
                parsed.append({
                    'type': 'paragraph',
                    'text': line
                })

            i += 1

        return parsed

    def _set_font(self, run, font_name, font_size, bold=False):
        """设置字体格式

        Args:
            run: docx的Run对象
            font_name: 字体名称
            font_size: 字号（磅值）
            bold: 是否加粗
        """
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold

        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_paragraph_with_format(self, text, style_type):
        """添加带格式的段落

        Args:
            text: 段落文本
            style_type: 样式类型（title/heading1/heading2/heading3/paragraph）
        """
        # 替换引号
        text = self._replace_quotes(text)

        # 创建段落
        para = self.doc.add_paragraph()

        # 设置段落格式
        para_format = para.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_format.line_spacing = Pt(28.9)  # 行距28.9磅
        para_format.space_after = Pt(0)  # 段后间距为0

        # 根据样式类型设置格式
        if style_type == 'title':
            # 文档标题：2号方正小标宋简体，居中，段后空35磅
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para_format.space_after = Pt(35)   # 段后间距35磅
            self._add_text_with_numbers(para, text,
                                       self.FONTS['方正小标宋简体'],
                                       self.FONT_SIZES['2号'],
                                       bold=False)

        elif style_type == 'heading1':
            # 一级标题：黑体三号，两端对齐，不加粗，首行缩进2字符
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_format.first_line_indent = Pt(32)  # 首行缩进2字符（三号字16磅×2=32磅）
            self._add_text_with_numbers(para, text,
                                       self.FONTS['黑体'],
                                       self.FONT_SIZES['3号'],
                                       bold=False)

        elif style_type == 'heading2':
            # 二级标题：楷体GB2312三号，两端对齐，首行缩进2字符
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_format.first_line_indent = Pt(32)  # 首行缩进2字符（三号字16磅×2=32磅）
            self._add_text_with_numbers(para, text,
                                       self.FONTS['楷体GB2312'],
                                       self.FONT_SIZES['3号'],
                                       bold=True)

        elif style_type == 'heading3':
            # 三级标题：仿宋GB2312三号，两端对齐，首行缩进2字符
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_format.first_line_indent = Pt(32)  # 首行缩进2字符（三号字16磅×2=32磅）
            self._add_text_with_numbers(para, text,
                                       self.FONTS['仿宋GB2312'],
                                       self.FONT_SIZES['3号'],
                                       bold=False)

        else:  # paragraph
            # 正文：仿宋GB2312三号，两端对齐，首行缩进2字符
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_format.first_line_indent = Pt(32)  # 首行缩进2字符（三号字16磅×2=32磅）
            self._add_text_with_bold(para, text,
                                    self.FONTS['仿宋GB2312'],
                                    self.FONT_SIZES['3号'])

    def _add_text_with_numbers(self, para, text, font_name, font_size, bold=False):
        """添加文本，数字和英文使用Times New Roman字体

        Args:
            para: 段落对象
            text: 文本内容
            font_name: 中文字体名称
            font_size: 字号
            bold: 是否加粗
        """
        # 使用正则表达式分割文本、数字和英文
        pattern = r'([a-zA-Z0-9]+)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            run = para.add_run(part)

            # 如果是数字或英文，使用Times New Roman
            if re.match(r'^[a-zA-Z0-9]+$', part):
                self._set_font(run, self.FONTS['Times New Roman'], font_size, bold)
            else:
                self._set_font(run, font_name, font_size, bold)

    def _add_text_with_bold(self, para, text, font_name, font_size):
        """添加文本，处理加粗标记

        Args:
            para: 段落对象
            text: 文本内容
            font_name: 字体名称
            font_size: 字号
        """
        # 使用正则表达式匹配加粗文本 **text**
        pattern = r'\*\*(.+?)\*\*'
        last_end = 0

        for match in re.finditer(pattern, text):
            # 添加加粗前的普通文本
            if match.start() > last_end:
                normal_text = text[last_end:match.start()]
                self._add_text_with_numbers(para, normal_text, font_name, font_size, bold=False)

            # 添加加粗文本
            bold_text = match.group(1)
            self._add_text_with_numbers(para, bold_text, font_name, font_size, bold=True)

            last_end = match.end()

        # 添加剩余的普通文本
        if last_end < len(text):
            normal_text = text[last_end:]
            self._add_text_with_numbers(para, normal_text, font_name, font_size, bold=False)

    def _create_custom_bullet_style(self):
        """创建自定义项目符号样式（使用✧符号）"""
        # 获取文档的样式部分
        styles = self.doc.styles

        # 尝试获取或创建自定义列表样式
        style_name = 'CustomBulletList'
        if style_name in [s.name for s in styles]:
            return styles[style_name]

        # 创建新样式
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['List Bullet']

        return style

    def _set_paragraph_bullet(self, para, level=0):
        """设置段落的项目符号为✧

        Args:
            para: 段落对象
            level: 列表层级（0-8）
        """
        # 确保段落有编号属性
        pPr = para._element.get_or_add_pPr()

        # 检查是否已有 numbering
        if pPr.numPr is None:
            # 创建 numbering 引用
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)

            # 设置编号ID（使用1）
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(min(level, 8)))

            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')

            numPr.append(ilvl)
            numPr.append(numId)

        # 添加自定义项目符号
        # 这需要更复杂的 OXML 操作，这里简化处理
        # 直接在段落开头添加 ✧ 符号作为替代
        pass

    def _add_list_item(self, text, level=0):
        """添加列表项段落

        Args:
            text: 列表项文本
            level: 列表层级（0-based）
        """
        # 替换引号
        text = self._replace_quotes(text)

        # 创建段落
        para = self.doc.add_paragraph()

        # 设置段落格式
        para_format = para.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_format.line_spacing = Pt(28.9)
        para_format.space_after = Pt(0)

        # 设置首行缩进和悬挂缩进
        indent_pt = 32 + level * 32  # 每级缩进增加32磅
        hanging_pt = 32  # 悬挂缩进32磅，让项目符号突出

        para_format.first_line_indent = Pt(-hanging_pt)  # 负的悬挂缩进
        para_format.left_indent = Pt(indent_pt + hanging_pt)

        # 添加项目符号和文本
        bullet_run = para.add_run('✧ ')
        self._set_font(bullet_run, self.FONTS['仿宋GB2312'], self.FONT_SIZES['3号'], bold=False)

        # 添加文本内容
        self._add_text_with_bold(para, text, self.FONTS['仿宋GB2312'], self.FONT_SIZES['3号'])

    def _add_table(self, headers, rows):
        """添加表格

        Args:
            headers: 表头列表
            rows: 数据行列表（每行是一个单元格列表）
        """
        if not headers or not rows:
            return

        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header

        # 创建表格
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'  # 使用带边框的样式

        # 设置表头
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text

            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # 设置表头单元格格式
            for paragraph in cell.paragraphs:
                # 水平居中
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    self._set_font(run, self.FONTS['仿宋GB2312'], self.FONT_SIZES['3号'], bold=True)

            # 设置表头底纹为浅灰色
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')  # 浅灰色 RGB(217,217,217)
            cell._element.get_or_add_tcPr().append(shading_elm)

        # 设置表格内容
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row_cells[col_idx]
                    cell.text = cell_text

                    # 垂直居中
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    # 设置单元格格式
                    for paragraph in cell.paragraphs:
                        # 水平居中
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            self._set_font(run, self.FONTS['仿宋GB2312'], self.FONT_SIZES['3号'], bold=False)

    def _convert_to_pdf(self):
        """将DOCX转换为PDF

        使用win32com调用Word进行转换（Windows系统）
        """
        try:
            import win32com.client
            import pythoncom

            print(f"正在转换为PDF: {self.output_pdf}")

            # 初始化COM
            pythoncom.CoInitialize()

            # 创建Word应用对象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # 打开DOCX文件
            doc = word.Documents.Open(str(self.output_docx.absolute()))

            # 保存为PDF (格式代码17表示PDF)
            doc.SaveAs(str(self.output_pdf.absolute()), FileFormat=17)

            # 关闭文档和Word
            doc.Close()
            word.Quit()

            # 清理COM
            pythoncom.CoUninitialize()

            print(f"PDF转换完成: {self.output_pdf}")

        except ImportError:
            print("警告: 未安装pywin32库，无法生成PDF")
            print("请运行: pip install pywin32")
        except Exception as e:
            print(f"警告: PDF转换失败: {e}")
            print("请确保系统已安装Microsoft Word")

    def convert(self):
        """执行转换"""
        if self.input_file:
            print(f"正在读取文件: {self.input_file}")

        # 读取markdown文件或内容
        if self.input_file:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            content = self.input_content

        # 解析markdown
        print("正在解析markdown内容...")
        parsed_data = self._parse_markdown(content)

        # 生成Word文档
        print("正在生成Word文档...")
        for item in parsed_data:
            if item['type'] == 'empty':
                continue
            elif item['type'] == 'list_item':
                self._add_list_item(item['text'], item.get('level', 0))
            elif item['type'] == 'table':
                self._add_table(item['headers'], item['rows'])
            else:
                self._add_paragraph_with_format(item['text'], item['type'])

        # 创建输出文件夹
        print(f"正在创建文件夹: {self.output_folder}")
        self.output_folder.mkdir(parents=True, exist_ok=True)

        # 保存Word文档
        print(f"正在保存Word文档: {self.output_docx}")
        self.doc.save(str(self.output_docx))
        print(f"Word文档保存完成: {self.output_docx}")

        # 转换为PDF
        self._convert_to_pdf()

        # 只有从文件读取时才移动原始MD文件
        if self.input_file:
            print(f"正在移动Markdown文件到: {self.output_md}")
            shutil.move(str(self.input_file), str(self.output_md))
            print(f"Markdown文件移动完成: {self.output_md}")
        else:
            # 从内容创建时，保存内容到MD文件
            print(f"正在保存Markdown文件到: {self.output_md}")
            with open(self.output_md, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"Markdown文件保存完成: {self.output_md}")

        print("\n" + "="*50)
        print("转换完成！")
        print(f"输出文件夹: {self.output_folder}")
        print(f"  - Markdown: {self.output_md.name}")
        print(f"  - Word: {self.output_docx.name}")
        if self.output_pdf.exists():
            print(f"  - PDF: {self.output_pdf.name}")
        print("="*50)


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <markdown文件路径>")
        print("示例: python md2docx.py 垂直领域大模型体系建设与推广应用的若干思考.md")
        sys.exit(1)

    input_file = sys.argv[1]

    try:
        converter = MarkdownToDocxConverter(input_file)
        converter.convert()
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
