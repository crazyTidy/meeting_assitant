"""Test script for DOCX generator."""
import sys
import io
from pathlib import Path

# Set UTF-8 encoding for output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from app.utils.docx_generator import DocxGenerator


def test_markdown_parsing():
    """Test various markdown formats to ensure they're properly stripped."""

    # Test content with various markdown formats
    test_content = """# 主标题

这是一段普通文本，包含**粗体**和*斜体*文字。

## 二级标题

### 三级标题

**粗体标题**

（一）中文数字标题

一、另一类标题

1. 第一项内容
2. 第二项内容
3. 第三项内容

- 无序列表项1
- 无序列表项2
- 无序列表项3

a. 字母列表项1
b. 字母列表项2

---

| 表头1 | 表头2 | 表头3 |
|-------|-------|-------|
| 内容1 | 内容2 | 内容3 |

---

这是`代码`文本。

这是带有~~删除线~~的文本。

这是带有[链接](https://example.com)的文本。

---

这是普通段落。
"""

    generator = DocxGenerator()

    # Test the parsing
    print("Testing markdown parsing...")
    print("=" * 50)

    generator.parse_markdown_content(test_content)

    # Check the document content
    print("\nDocument paragraphs:")
    print("=" * 50)
    for i, para in enumerate(generator.doc.paragraphs):
        text = para.text
        # Check for markdown markers
        issues = []
        if '#' in text and any(text.startswith(c) for c in ['#']):
            issues.append("Contains heading markers")
        if '**' in text or '__' in text:
            issues.append("Contains bold markers")
        if '*' in text or '_' in text:
            issues.append("Contains italic markers")
        if '`' in text:
            issues.append("Contains code markers")
        if '~~' in text:
            issues.append("Contains strikethrough markers")
        if '[' in text and '](' in text:
            issues.append("Contains link markers")

        status = "[FAIL]" if issues else "[PASS]"
        print(f"[{i}] {status}: {repr(text)}")
        if issues:
            for issue in issues:
                print(f"    - {issue}")

    # Save test file
    output_path = Path(__file__).parent / "test_output.docx"
    generator.doc.save(output_path)
    print(f"\n[PASS] Test document saved to: {output_path}")
    print("  Please open the file and verify:")
    print("  - No markdown symbols are visible")
    print("  - Headings use proper Chinese fonts (黑体, 楷体GB2312)")
    print("  - Body text uses 仿宋_GB2312")
    print("  - No **, *, __, #, etc. characters appear in the document")


def test_inline_text_stripping():
    """Test that inline markdown is properly stripped."""
    from app.utils.docx_generator import DocxGenerator
    from docx import Document

    generator = DocxGenerator()

    test_cases = [
        ("**粗体**", "粗体"),
        ("__粗体__", "粗体"),
        ("*斜体*", "斜体"),
        ("_斜体_", "斜体"),
        ("***粗斜体***", "粗斜体"),
        ("___粗斜体___", "粗斜体"),
        ("~~删除线~~", "删除线"),
        ("`代码`", "代码"),
        ("[链接文字](https://example.com)", "链接文字"),
        ("混合**粗体**和*斜体*", "混合粗体和斜体"),
        ("前缀**粗体**后缀", "前缀粗体后缀"),
    ]

    print("\nTesting inline markdown stripping:")
    print("=" * 50)

    all_passed = True
    for input_text, expected in test_cases:
        doc = Document()
        para = doc.add_paragraph()
        generator._parse_inline_text(para, input_text)
        result = para.text

        passed = result == expected
        status = "[PASS]" if passed else "[FAIL]"
        all_passed = all_passed and passed

        print(f"{status} Input: {repr(input_text)}")
        print(f"  Expected: {repr(expected)}")
        print(f"  Got:      {repr(result)}")
        if not passed:
            print(f"  MISMATCH!")
        print()

    return all_passed


if __name__ == "__main__":
    print("DOCX Generator Test")
    print("=" * 50)

    # Test inline text stripping
    inline_passed = test_inline_text_stripping()

    # Test full document generation
    test_markdown_parsing()

    print("\n" + "=" * 50)
    if inline_passed:
        print("[PASS] All inline tests passed!")
    else:
        print("[FAIL] Some inline tests failed!")
    print("\nPlease check test_output.docx for visual verification.")
